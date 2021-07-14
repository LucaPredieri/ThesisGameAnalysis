####
## Importing Libraries.
####

from docx.shared import Pt
from docx2pdf import convert
# import os
import docx


####
## Function to read from the word file.
####

def ReadDoc(filename):
    file = open(filename, 'r')
    f = file.readlines()
    newList = []
    for line in f:
        newList.append(line.strip())
    return newList


####
## Reading the file and filling arrays.
####

name = 'Luca'
surname = 'Predieri'
age = '21'

counter_1 = 0
counter_2 = 0

infos = ReadDoc('test.txt')

for x in range(len(infos)):
    if infos[x] == '0001':          # Hit by a car.
        counter_a += 1
    elif infos[x] == '0010':        # Hit by a train.
        counter_b += 1
    elif infos[x] == '0011':        # Player didn't wait for the train.
        counter_c += 1
    elif infos[x] == '0100':        # Player waited for the train.
        counter_d += 1
    elif infos[x] == '0101':        # Player crossed the street <2s
        counter_e += 1
    elif infos[x] == '0110':        # Player crossed the street <4s
        counter_f += 1
    elif infos[x] == '0111':        # Player crossed the street <10s
        counter_g += 1
    elif infos[x] == '1000':        # Player crossed the street cutting the zebras.
        counter_h += 1
    elif infos[x] == '1001':        # Player crossed the street (no tl) well.
        counter_i += 1
    elif infos[x] == '1010':        # Player crossed the street (no tl) badly.
        counter_j += 1
    elif infos[x] == '1011':        # Player hit a person.
        counter_k += 1
    elif infos[x] == '1100':        # Player helps (?)
        counter_l += 1
    elif infos[x] == '1101':        # Player arrives.
        counter_m += 1
    elif infos[x] == '1110':        # Player walks in a bad way.
        counter_n += 1
    elif infos[x] == '1111':        # Player walks in a good way.
        counter_o += 1

####
## Creating word file.
####

doc = docx.Document()

## Title and introduction.
doc.add_heading("ExeRoads!", 0)
paragraph_a = doc.add_paragraph().add_run(
    'ExeRoads is a game made by Fabio Conti and Luca Predieri. '
    'The aim of the game is to get to the church paying '
    'attention to the obstacles and whatever is around the subject.')
paragraph_a.italic = True

## Player's info.
doc.add_heading("Player's info:", 1)
paragraph = doc.add_paragraph().add_run('')
paragraph.font.size = Pt(9)
paragraph = doc.add_paragraph('Name:   ')
paragraph.add_run(name).bold = True
paragraph = doc.add_paragraph('Last name:   ')
paragraph.add_run(surname).bold = True
paragraph = doc.add_paragraph('Age:   ')
paragraph.add_run(age).bold = True

## Player's score.
doc.add_heading("Player's score:", 1)

# Sidewalk related score section.
doc.add_heading("Road's related score:", 2)
ps = para.add_run(
    " The sidewalk score is based on how many crosswalk blocks the user collides."
)
if 0 < counter_1 < 20:
    para = doc.add_paragraph(
        "The user doesn't use the sidewalk as he should, only ")
    para.add_run(str(counter_i))
    para.add_run(
        ' points related to the sidewalk section are scored.')
elif 20 < counter_1 < 40:
    para = doc.add_paragraph(
        "The user uses the sidewalk as he should, ")
    para.add_run(str(counter_i))
    para.add_run(
        " are the points related to the sidewalk section scored.")

# Cars related score section.

# Traffic light related score section.

# People related score section.

## Saving the doc and converting into a PDF.
doc.save('prove.docx')
convert("prove.docx")
